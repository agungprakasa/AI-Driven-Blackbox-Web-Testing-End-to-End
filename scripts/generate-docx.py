#!/usr/bin/env python3
"""Generator laporan DOCX — mendukung API Testing dan Web Testing.

Usage:
    python scripts/generate-docx.py --mode api
    python scripts/generate-docx.py --mode web
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "reports" / "docx"


def baca(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def load_screenshot(tc_id: str, mode: str) -> Path | None:
    """Load screenshot for a test case. Matches any PNG with TC-ID prefix."""
    base = "evidence-web" if mode == "web" else "evidence"
    for folder in [ROOT / base / "PASS", ROOT / base / "FAIL", ROOT / base]:
        if folder.exists():
            # Match: TC-P001_screenshot.png OR TC-P001_xxx.png
            for f in folder.glob(f"{tc_id}*.png"):
                return f
    return None


def bersihkan_md(s: str) -> str:
    return re.sub(r"<br\s*/?>", " ", re.sub(r"`([^`]*)`", r"\1", s))


def set_cell_shading(cell, color_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def parse_tests(tc_teks: str, mode: str) -> list[dict]:
    tests = []
    for line in tc_teks.splitlines():
        s = line.strip()
        if s.startswith("|") and re.match(r"^\|\s*TC-[A-Z0-9]+-?\d{3}\b", s):
            cols = [c.strip() for c in s.strip("|").split("|")]
            if len(cols) < 7:
                continue
            m = re.match(r"TC-([A-Z0-9]+)-?", cols[0])
            if mode == "web":
                tests.append({
                    "id": cols[0], "judul": cols[1],
                    "halaman": cols[2], "aksi": cols[3],
                    "expected": cols[4], "priority": cols[5],
                    "status": cols[6].upper(),
                    "modul": m.group(1) if m else "UNKNOWN",
                })
            else:
                tests.append({
                    "id": cols[0], "judul": cols[1],
                    "endpoint": cols[2], "request": cols[3],
                    "expected": cols[4], "priority": cols[5],
                    "status": cols[6].upper(),
                    "modul": m.group(1) if m else "UNKNOWN",
                })
    return tests


def get_category(tc_id: str) -> str:
    prefix_map = {
        "P": "Normal Case", "N": "Abnormal Case", "B": "Boundary Value",
        "S": "Security", "PF": "Performance", "DV": "Data Validation",
        "E2E": "Integration / E2E", "W": "Web Functional",
        "WA": "Web Accessibility", "WR": "Web Responsive",
    }
    for prefix, cat in prefix_map.items():
        if re.match(rf"^TC-{prefix}\d", tc_id):
            return cat
    return "Lainnya"


def add_header(doc, mode: str, title: str, modul: str):
    """Add official company header to document."""
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.2)

    logo_path = ROOT / "config" / "logo-pos-indonesia.png"
    header = section.header
    header.is_linked_to_previous = False

    header_table = header.add_table(rows=4, cols=4, width=Inches(6.5))
    header_table.style = "Table Grid"
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_col_widths = [Cm(2.2), Cm(8.5), Cm(2.2), Cm(2.8)]
    for row in header_table.rows:
        for i, w in enumerate(hdr_col_widths):
            row.cells[i].width = w

    # Merge col 0 vertically (logo)
    header_table.rows[0].cells[0].merge(header_table.rows[3].cells[0])
    # Merge col 1 row 0..1 (company name)
    header_table.rows[0].cells[1].merge(header_table.rows[1].cells[1])

    # Logo
    cell_logo = header_table.rows[0].cells[0]
    cell_logo.text = ""
    if logo_path.exists():
        p = cell_logo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo_path), width=Cm(2.0))
    tc = cell_logo._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)

    # Company name
    cell_company = header_table.rows[0].cells[1]
    cell_company.text = ""
    p = cell_company.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PT POS INDONESIA (PERSERO)")
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    p2 = cell_company.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("BAG. QUALITY ASSURANCE DEVSECOPS PLATFORM")
    run2.font.bold = True
    run2.font.size = Pt(9)
    run2.font.name = "Calibri"
    tc_c = cell_company._tc
    tcPr_c = tc_c.get_or_add_tcPr()
    vAlign_c = OxmlElement("w:vAlign")
    vAlign_c.set(qn("w:val"), "center")
    tcPr_c.append(vAlign_c)

    # Row 0-1 col 2-3
    set_cell_text(header_table.rows[0].cells[2], "Dokumen", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header_table.rows[0].cells[3], "Skenario", bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header_table.rows[1].cells[2], "Halaman", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header_table.rows[1].cells[3], "", size=9)

    # Row 2-3
    set_cell_text(header_table.rows[2].cells[1], title, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header_table.rows[2].cells[2], "No. Revisi", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header_table.rows[2].cells[3], "-", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header_table.rows[3].cells[1], f"Modul : {modul}", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header_table.rows[3].cells[2], "Tanggal", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header_table.rows[3].cells[3], date.today().strftime("%d/%m/%Y"), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate DOCX report")
    parser.add_argument("--mode", choices=["api", "web"], default="web", help="Testing mode: api or web")
    args = parser.parse_args()
    mode = args.mode

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    tc_file = "test-case-web.md" if mode == "web" else "test-case.md"
    defect_file = "defect-list-web.md" if mode == "web" else "defect-list.md"

    tc_teks = baca(ROOT / "docs" / tc_file)
    defect_teks = baca(ROOT / "docs" / defect_file)

    tests = parse_tests(tc_teks, mode)
    if not tests:
        print(f"ERROR: Tidak ada test case ditemukan di docs/{tc_file}")
        return 1

    total = len(tests)
    hitung = Counter(t["status"] for t in tests)
    n_pass = hitung.get("PASS", 0)
    n_fail = hitung.get("FAIL", 0)
    pass_rate = f"{n_pass / total * 100:.1f}%" if total else "-"

    # Mode-specific config
    if mode == "web":
        doc_title = "PENGUJIAN WEB"
        modul_name = "Web Application"
        subject = "AI-Driven Blackbox Web Testing"
        note = "Pengujian fungsionalitas, keamanan, dan responsivitas halaman web"
        cred_lines = [
            "A. Base URL",
            "1. URL : (sesuaikan dengan URL web yang diuji)",
            "2. Browser : Chromium (Playwright)",
            "B. Authentication : (sesuaikan)",
        ]
        summary_desc = (
            f"Pengujian Web dilakukan terhadap {total} test case dengan pass rate {pass_rate}. "
            f"Pengujian mencakup functional testing, accessibility, responsive design, "
            f"security, performance, dan end-to-end flow."
        )
        table_title = "WEB APPLICATION"
        col_headers = ["No.", "ID", "Butir\nUji", "Halaman/\nURL", "Aksi/\nInput",
                       "Hasil\nPengujian", "Status\nPengujian", "Keterangan", "Evidence"]
        col_widths = [Cm(0.8), Cm(1.5), Cm(2.2), Cm(3.0), Cm(4.0), Cm(3.5), Cm(1.8), Cm(3.0), Cm(6.0)]
        num_cols = 9
    else:
        doc_title = "PENGUJIAN API TARIFF"
        modul_name = "Tariff (getfeeLnDiscountNew)"
        subject = "API Tariff — getfeeLnDiscountNew"
        note = "AI-Driven Blackbox API Testing — End-to-End"
        cred_lines = [
            "A. Base URL",
            "1. URL : http://10.29.41.37:8280/test/1.0.0",
            "2. Endpoint : POST /getfeeLnDiscountNew",
            "B. Authentication : None (tanpa autentikasi)",
        ]
        summary_desc = (
            f"Pengujian API dilakukan terhadap {total} test case dengan pass rate {pass_rate}. "
            f"Endpoint yang diuji adalah POST /getfeeLnDiscountNew pada base URL "
            f"http://10.29.41.37:8280/test/1.0.0. Pengujian mencakup normal case, "
            f"abnormal case, boundary value, security, performance, data validation, "
            f"dan integrasi end-to-end."
        )
        table_title = "API TARIFF — getfeeLnDiscountNew"
        col_headers = ["No.", "ID", "Butir\nUji", "Uraian\nKegiatan",
                       "Hasil\nPengujian", "Status\nPengujian", "Keterangan", "Evidence"]
        col_widths = [Cm(1.0), Cm(1.8), Cm(2.5), Cm(5.5), Cm(4.0), Cm(2.0), Cm(3.5), Cm(6.0)]
        num_cols = 8

    doc = Document()

    # =========================================================
    # HEADER
    # =========================================================
    add_header(doc, mode, doc_title, modul_name)

    # =========================================================
    # BODY CONTENT
    # =========================================================
    header_data = [
        ("Tanggal\t\t:", date.today().strftime("%d/%m/%Y")),
        ("Penguji\t\t\t:", "Agung Prakasa"),
        ("Tempat\t\t\t:", "Graha Pos Indonesia, Lt. 4"),
        ("Subjek\t\t\t:", subject),
        ("Note\t\t\t:", note),
    ]
    for label, value in header_data:
        p = doc.add_paragraph()
        run = p.add_run(f"{label} {value}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_paragraph()

    # Kredensial
    p = doc.add_paragraph()
    run = p.add_run("Data Akses Kredensial Pengujian:")
    run.font.bold = True
    run.font.size = Pt(11)
    for line in cred_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    doc.add_paragraph()

    # Keterangan
    p = doc.add_paragraph()
    run = p.add_run("Keterangan :")
    run.font.bold = True
    run.font.size = Pt(11)

    keterangan_table = doc.add_table(rows=7, cols=2)
    keterangan_table.style = "Table Grid"
    keterangan_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in keterangan_table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(9.5)

    keterangan_data = [
        ("Status Pengujian", ""),
        ("OK", "Jika hasil pengujian diterima"),
        ("NOK", "Jika hasil pengujian tidak diterima"),
        ("Keterangan", ""),
        ("P", "Perbaikan"),
        ("T", "Tambahan / Usulan"),
        ("Status Perbaikan", "OK = Diterima, NOK = Tidak diterima"),
    ]
    for ri, (label, desc) in enumerate(keterangan_data):
        is_header = desc == ""
        set_cell_text(keterangan_table.rows[ri].cells[0], label, bold=is_header, size=9)
        set_cell_text(keterangan_table.rows[ri].cells[1], desc, size=9)
        if is_header:
            set_cell_shading(keterangan_table.rows[ri].cells[0], "D6E4F0")
            set_cell_shading(keterangan_table.rows[ri].cells[1], "D6E4F0")

    doc.add_paragraph()

    # =========================================================
    # RINGKASAN EKSEKUTIF
    # =========================================================
    p = doc.add_paragraph()
    run = p.add_run("Ringkasan Eksekutif")
    run.font.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(summary_desc)
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    doc.add_paragraph()

    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in summary_table.rows:
        row.cells[0].width = Cm(6.0)
        row.cells[1].width = Cm(4.0)

    set_cell_text(summary_table.rows[0].cells[0], "Metrik", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(summary_table.rows[0].cells[1], "Nilai", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(summary_table.rows[0].cells[0], "4472C4")
    set_cell_shading(summary_table.rows[0].cells[1], "4472C4")
    for c in [0, 1]:
        for p in summary_table.rows[0].cells[c].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    defect_count = len([l for l in defect_teks.splitlines()
                        if l.strip().startswith("|") and re.match(r"^\|\s*DEF-\d+", l.strip())])
    summary_data = [
        ("Total Test Case", str(total)),
        ("PASS", str(n_pass)),
        ("FAIL", str(n_fail)),
        ("Pass Rate", pass_rate),
        ("Jenis Pengujian", "7 kategori"),
        ("Total Defect", str(defect_count)),
    ]
    for label, value in summary_data:
        row = summary_table.add_row()
        set_cell_text(row.cells[0], label, size=9)
        set_cell_text(row.cells[1], value, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # =========================================================
    # LANDSCAPE SECTION FOR MAIN TABLE
    # =========================================================
    new_section = doc.add_section()
    new_section.orientation = 1
    new_section.page_width = Cm(29.7)
    new_section.page_height = Cm(21.0)
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)
    new_section.top_margin = Cm(1.5)
    new_section.bottom_margin = Cm(1.5)

    # =========================================================
    # MAIN TABLE
    # =========================================================
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for i, h in enumerate(col_headers):
        set_cell_text(hdr_row.cells[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr_row.cells[i], "4472C4")
        for p in hdr_row.cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Title row (merged)
    row_title = table.add_row()
    set_cell_text(row_title.cells[0], table_title, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(row_title.cells[0], "D6E4F0")
    for i in range(1, num_cols):
        set_cell_shading(row_title.cells[i], "D6E4F0")
    for i in range(1, num_cols):
        row_title.cells[0].merge(row_title.cells[i])

    # Group by category
    categories = []
    current_cat = None
    for t in tests:
        cat = get_category(t["id"])
        if cat != current_cat:
            categories.append({"name": cat, "tests": []})
            current_cat = cat
        categories[-1]["tests"].append(t)

    no_counter = 0
    for cat_info in categories:
        cat_name = cat_info["name"]
        cat_tests = cat_info["tests"]

        row_cat = table.add_row()
        set_cell_text(row_cat.cells[0], cat_name, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(row_cat.cells[0], "E2EFDA")
        for i in range(1, num_cols):
            set_cell_shading(row_cat.cells[i], "E2EFDA")
        for i in range(1, num_cols):
            row_cat.cells[0].merge(row_cat.cells[i])

        for t in cat_tests:
            no_counter += 1
            tc_id = t["id"]
            status = t["status"]
            ok_nok = "OK" if status == "PASS" else "NOK"
            screenshot = load_screenshot(tc_id, mode)

            row = table.add_row()

            # No.
            set_cell_text(row.cells[0], str(no_counter), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            # ID
            set_cell_text(row.cells[1], tc_id, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            # Butir Uji
            set_cell_text(row.cells[2], t["judul"], size=8)

            if mode == "web":
                # Halaman/URL
                set_cell_text(row.cells[3], t.get("halaman", ""), size=8)
                # Aksi/Input
                set_cell_text(row.cells[4], t.get("aksi", ""), size=8)
                # Hasil Pengujian
                set_cell_text(row.cells[5], f"Status: {status}", size=8)
                # Status
                set_cell_text(row.cells[6], ok_nok, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
                if ok_nok == "OK":
                    set_cell_shading(row.cells[6], "C6EFCE")
                else:
                    set_cell_shading(row.cells[6], "FFC7CE")
                # Keterangan
                keterangan = ""
                if status == "FAIL":
                    keterangan = "P: Perlu perbaikan"
                set_cell_text(row.cells[7], keterangan, size=8)
                # Evidence
                ev_col = 8
            else:
                # Uraian Kegiatan
                deskripsi = f"Request: {t['request']}\nExpected: {t['expected']}"
                set_cell_text(row.cells[3], deskripsi, size=8)
                # Hasil Pengujian
                set_cell_text(row.cells[4], f"Status: {status}", size=8)
                # Status
                set_cell_text(row.cells[5], ok_nok, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
                if ok_nok == "OK":
                    set_cell_shading(row.cells[5], "C6EFCE")
                else:
                    set_cell_shading(row.cells[5], "FFC7CE")
                # Keterangan
                keterangan = ""
                if status == "FAIL":
                    if "S" in tc_id[:4]:
                        keterangan = "P: Perlu input validation"
                    elif "N" in tc_id[:4]:
                        keterangan = "P: Perlu validasi input"
                    elif "B" in tc_id[:4]:
                        keterangan = "P: Perlu boundary check"
                    else:
                        keterangan = "P: Perlu perbaikan"
                set_cell_text(row.cells[6], keterangan, size=8)
                ev_col = 7

            # Evidence (screenshot)
            if screenshot and screenshot.exists():
                cell_ev = row.cells[ev_col]
                cell_ev.text = ""
                p = cell_ev.paragraphs[0]
                run = p.add_run()
                run.add_picture(str(screenshot), width=Inches(2.2))
            else:
                set_cell_text(row.cells[ev_col], "-", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in table.rows:
        for i, w in enumerate(col_widths):
            if i < len(row.cells):
                row.cells[i].width = w

    # =========================================================
    # BACK TO PORTRAIT
    # =========================================================
    new_section2 = doc.add_section()
    new_section2.orientation = 0
    new_section2.page_width = Cm(21.0)
    new_section2.page_height = Cm(29.7)
    new_section2.left_margin = Cm(2.5)
    new_section2.right_margin = Cm(2.5)
    new_section2.top_margin = Cm(2.5)
    new_section2.bottom_margin = Cm(2.5)

    # =========================================================
    # DAFTAR DEFECT
    # =========================================================
    doc.add_paragraph()
    doc.add_heading("Daftar Defect", level=1)

    defect_lines = [l.strip() for l in defect_teks.splitlines()
                    if l.strip().startswith("|") and re.match(r"^\|\s*DEF-\d+", l.strip())]
    if defect_lines:
        defect_table = doc.add_table(rows=1, cols=5)
        defect_table.style = "Table Grid"
        defect_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        d_headers = ["ID", "Judul", "Severity", "Status", "Keterangan"]
        for i, h in enumerate(d_headers):
            set_cell_text(defect_table.rows[0].cells[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(defect_table.rows[0].cells[i], "4472C4")
            for p in defect_table.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for line in defect_lines:
            cols = [c.strip() for c in line.strip("|").split("|")]
            if len(cols) >= 4:
                row = defect_table.add_row()
                for i in range(min(5, len(cols))):
                    set_cell_text(row.cells[i], cols[i], size=8)
    else:
        doc.add_paragraph("Tidak ada defect terdokumentasi.")

    # =========================================================
    # KESIMPULAN & REKOMENDASI
    # =========================================================
    doc.add_heading("Kesimpulan & Rekomendasi", level=1)
    doc.add_paragraph(
        f"Dari {total} test case yang dieksekusi, {n_pass} PASS dan {n_fail} FAIL "
        f"dengan pass rate {pass_rate}."
    )
    p = doc.add_paragraph()
    run = p.add_run("Rekomendasi:")
    run.font.bold = True

    if n_fail > 0:
        if mode == "web":
            doc.add_paragraph("1. Perbaiki elemen UI yang tidak responsif pada berbagai ukuran layar.", style="List Number")
            doc.add_paragraph("2. Perkuat validasi form input — pastikan semua input tervalidasi di sisi client dan server.", style="List Number")
            doc.add_paragraph("3. Tambahkan accessibility attributes (ARIA labels, alt text) pada elemen interaktif.", style="List Number")
        else:
            doc.add_paragraph("1. Perbaiki validasi input pada endpoint — beberapa input invalid masih diterima dengan status 200.", style="List Number")
            doc.add_paragraph("2. Perkuat proteksi keamanan — SQL Injection, XSS, dan Path Traversal menyebabkan HTTP 500.", style="List Number")
            doc.add_paragraph("3. Tambahkan validasi Content-Type dan request method.", style="List Number")
    else:
        doc.add_paragraph("Semua test case PASS. Tidak ada rekomendasi perbaikan saat ini.")

    # =========================================================
    # SAVE
    # =========================================================
    mode_label = "Web" if mode == "web" else "API"
    # Get endpoint name from test case
    endpoint_name = "getfeeLnDiscountNew"
    out = OUT_DIR / f"Laporan-{endpoint_name}-{date.today().isoformat()}.docx"
    doc.save(out)
    print(f"Laporan DOCX dibuat: {out.relative_to(ROOT)}")
    print(f"Mode: {mode.upper()} | Total: {total} | PASS: {n_pass} | FAIL: {n_fail} | Rate: {pass_rate}")

    ss_count = 0
    base = "evidence-web" if mode == "web" else "evidence"
    for folder in [ROOT / base / "PASS", ROOT / base / "FAIL", ROOT / base]:
        if folder.exists():
            ss_count += len(list(folder.glob("*_screenshot.png")))
    print(f"Screenshots terlampir: {ss_count}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
